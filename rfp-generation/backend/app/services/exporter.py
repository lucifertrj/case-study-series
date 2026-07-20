from __future__ import annotations

import io
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt


def build_docx(project: dict[str, Any], include_drafts: bool = False) -> io.BytesIO:
    doc = DocxDocument()

    title = doc.add_heading(f"{project['name']} — RFP Response", level=0)
    for run in title.runs:
        run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.add_run(f"Prepared for: {project.get('client', '')}").italic = True
    if project.get("due_date"):
        sub.add_run(f"\nDue date: {project['due_date']}").italic = True

    doc.add_paragraph()

    questions = sorted(project.get("questions", []), key=lambda q: q.get("number", 0))
    exported = 0

    for q in questions:
        st = q.get("status")
        answer = q.get("answer")
        if not answer or not answer.get("content"):
            continue
        if st == "approved" or (include_drafts and st == "draft"):
            heading = doc.add_heading(f"Q{q['number']}. {q['text']}", level=2)
            for run in heading.runs:
                run.font.size = Pt(13)

            doc.add_paragraph(answer["content"])

            names = list({
                s.get("document_name")
                for s in answer.get("sources", [])
                if s.get("document_name")
            })
            if names:
                foot = doc.add_paragraph()
                r = foot.add_run(f"Sources: {', '.join(sorted(names))}")
                r.italic = True
                r.font.size = Pt(9)

            doc.add_paragraph()
            exported += 1

    if exported == 0:
        doc.add_paragraph("No approved answers to export yet.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
